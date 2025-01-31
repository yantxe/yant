import sys
import pandas as pd
from PyQt5.QtCore import QTimer, QTime
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QLabel, QTableWidget, QTableWidgetItem, QHBoxLayout, QFileDialog, QComboBox
from PyQt5.QtGui import QFont
from docx import Document
import os

class StopwatchApp(QWidget):
    def __init__(self):
        super().__init__()

        # Initialize UI components
        self.timer = QTimer(self)
        self.time = QTime(0, 0, 0)  # Current stopwatch time
        self.start_time = None  # Store the start time of the stopwatch
        self.last_lap_time = None  # To store the last lap time
        self.laps = []

        self.initUI()

    def initUI(self):
        self.setWindowTitle('Stopwatch App')

        # Set font size for the entire app
        font = QFont()
        font.setPointSize(30)

        # Display the timer
        self.timer.timeout.connect(self.updateTime)
        self.time_display = QLabel(self.time.toString("mm:ss.zz"))
        self.time_display.setFont(font)
        
        # Buttons
        self.start_button = QPushButton('Start', self)
        self.stop_button = QPushButton('Stop', self)
        self.lap_button = QPushButton('Lap', self)
        self.reset_button = QPushButton('Reset', self)
        self.export_button = QPushButton('Export to Word', self)
        
        self.start_button.setFont(font)
        self.stop_button.setFont(font)
        self.lap_button.setFont(font)
        self.reset_button.setFont(font)
        self.export_button.setFont(font)
        
        self.start_button.clicked.connect(self.start)
        self.stop_button.clicked.connect(self.stop)
        self.lap_button.clicked.connect(self.lap)
        self.reset_button.clicked.connect(self.reset)
        self.export_button.clicked.connect(self.export_data)

        # Dropdown for selecting the name
        self.name_select = QComboBox(self)
        self.name_select.addItems(['Danang', 'Tama', 'Farhan', 'Candra'])
        self.name_select.setFont(font)
        
        # Layout
        layout = QVBoxLayout()
        layout.addWidget(self.time_display)
        
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.start_button)
        button_layout.addWidget(self.stop_button)
        button_layout.addWidget(self.lap_button)
        button_layout.addWidget(self.reset_button)
        layout.addLayout(button_layout)
        
        layout.addWidget(self.name_select)  # Add name selection dropdown
        
        # Table for lap data
        self.table = QTableWidget(self)
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(['Lap #', 'Lap Time', 'Total Time'])
        self.table.setFont(font)  # Apply font size to table

        # Adjust column width
        self.table.setColumnWidth(0, 200)  # Column width for the first column
        self.table.setColumnWidth(1, 400)  # Column width for the second column
        self.table.setColumnWidth(2, 400)  # Column width for the third column

        layout.addWidget(self.table)
        
        layout.addWidget(self.export_button)
        
        self.setLayout(layout)

        # Resize window to 1200 x 1200
        self.resize(1200, 1200)
        self.show()

    def start(self):
        self.timer.start(100)  # 100 ms interval
        self.start_time = self.time  # Store the start time when stopwatch starts
        self.last_lap_time = self.time  # Store the first lap time
        self.start_button.setEnabled(False)
        self.stop_button.setEnabled(True)

    def stop(self):
        self.timer.stop()
        self.start_button.setEnabled(True)
        self.stop_button.setEnabled(False)

    def lap(self):
        # Lap time shows the time when the lap button is pressed
        lap_time = self.time.toString("mm:ss.zz")

        # Calculate the total time by measuring the difference between the first lap and the current lap
        total_time = self.last_lap_time.msecsTo(self.time)

        # Append lap data (Lap number, Lap time, Total time)
        self.laps.append([len(self.laps) + 1, lap_time, self.formatTime(total_time)])

        # Update table and scroll to the latest lap
        self.updateTable()
        self.scroll_to_latest_lap()

        # Set the current lap time as the last lap time for the next lap
        self.last_lap_time = self.time

    def reset(self):
        self.timer.stop()
        self.time = QTime(0, 0, 0)
        self.start_time = None
        self.last_lap_time = None
        self.laps.clear()
        self.updateTable()
        self.start_button.setEnabled(True)
        self.stop_button.setEnabled(False)

    def updateTime(self):
        self.time = self.time.addMSecs(100)
        self.time_display.setText(self.time.toString("mm:ss.zz"))

    def updateTable(self):
        self.table.setRowCount(len(self.laps))
        
        # Set row height dynamically based on font size and row count
        row_height = 60  # Increase row height if needed
        for i in range(len(self.laps)):
            self.table.setRowHeight(i, row_height)
        
        # Update each row with lap data
        for i, lap in enumerate(self.laps):
            self.table.setItem(i, 0, QTableWidgetItem(str(lap[0])))
            self.table.setItem(i, 1, QTableWidgetItem(lap[1]))
            self.table.setItem(i, 2, QTableWidgetItem(lap[2]))

    def scroll_to_latest_lap(self):
        # Scroll to the last item in the table (latest lap)
        self.table.scrollToItem(self.table.item(len(self.laps) - 1, 0), QTableWidget.PositionAtBottom)

    def export_data(self):
        # Get selected name
        selected_name = self.name_select.currentText()

        # Define the file name based on the selected person
        file_name = f"{selected_name}_lap_times.docx"
        
        # Check if the file exists and read it; if it exists, remove the existing data
        if os.path.exists(file_name):
            document = Document(file_name)  # Open existing Word file
            # Clear the existing paragraphs (previous test data)
            document.paragraphs.clear()
        else:
            document = Document()  # Create new Word document

        # Get the next available test number
        test_number = len([p for p in document.paragraphs if p.text.startswith('Test')]) + 1
        document.add_paragraph(f"Test {test_number}")  # Add Test number title

        # Add lap data with consistent spacing
        for lap in self.laps:
            # Format the lap data with fixed width for each column
            lap_data = f"{str(lap[0]).zfill(2)}      {lap[1]}      {lap[2]}"
            document.add_paragraph(lap_data)

        # Save the document
        document.save(file_name)

    def formatTime(self, ms):
        """Format time in milliseconds to mm:ss.xx"""
        return QTime(0, 0).addMSecs(ms).toString("mm:ss.zz")
    
if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = StopwatchApp()
    sys.exit(app.exec_())
